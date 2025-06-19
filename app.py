import streamlit as st
import pandas as pd
from supabase import create_client
import time
from docx import Document
from io import BytesIO
import os
from docx.shared import RGBColor, Pt
import openpyxl
import datetime

url: str = "https://etxhbhpjqoaoowfoscob.supabase.co"
key: str = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImV0eGhiaHBqcW9hb293Zm9zY29iIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc0NDcwNTAzNiwiZXhwIjoyMDYwMjgxMDM2fQ.xJ7aC9TrQU-qbdf6KBB9D4FmyANf3NSvEHhpO4-lUEQ"

table_name ="REFERENCEDATABASEN"

#######################################################
####################### Funktioner ####################
#######################################################

# OPRET FORBINDELSE TIL SUPABASE
@st.cache_resource
def init_supabase():
    return create_client(url, key)

# HENT DATA FRA SQL
def fetch_data(supabase, table_name, batch_size=1000):
    all_data = []
    offset = 0

    while True:
        response = supabase.table(table_name).select("*").range(offset, offset + batch_size - 1).execute()
        data = response.data

        if not data:
            break

        all_data.extend(data)
        offset += batch_size

    df = pd.DataFrame(all_data)
    df = df.sort_values('opgave_id', ascending=False)
    return df

# OPDATER RÆKKER I SQL FRA STREAMLIT
def update_multiple_rows(supabase, table_name, edited_data, original_data):
    try:
        # Check if key column exists
        if 'opgave_id' not in edited_data.columns or 'opgave_id' not in original_data.columns:
            return False, "Kolonnen 'opgave_id' mangler i data."

        edited_df = edited_data.copy()

        # Remove UI-specific columns like 'Select'
        edited_df = edited_df.drop(columns=['Select'], errors='ignore')

        # Align datetime formats
        datetime_cols = [
            col for col in edited_df.columns
            if pd.api.types.is_datetime64_any_dtype(edited_df[col]) or 
               pd.api.types.is_datetime64_any_dtype(original_data[col])
        ]

        for col in datetime_cols:
            edited_df[col] = pd.to_datetime(edited_df[col], errors='coerce')
            original_data[col] = pd.to_datetime(original_data[col], errors='coerce')

        changes = []

        # Compare rows
        for _, row in edited_df.iterrows():
            opgave_id = row['opgave_id']
            original_row = original_data[original_data['opgave_id'] == opgave_id]

            if original_row.empty:
                continue  # Possibly a new row, ignore

            original_dict = original_row.iloc[0].to_dict()
            row_dict = row.to_dict()

            # Detect changes
            has_changed = any(
                (pd.isna(row_dict[k]) != pd.isna(original_dict.get(k))) or 
                (not pd.isna(row_dict[k]) and row_dict[k] != original_dict.get(k))
                for k in row_dict if k in original_dict
            )

            if has_changed:
                update_dict = {
                    k: (
                        v.isoformat() if isinstance(v, (pd.Timestamp, pd.DatetimeTZDtype)) else
                        None if pd.isna(v) else v
                    )
                    for k, v in row_dict.items()
                }

                changes.append({'opgave_id': opgave_id, 'data': update_dict})

        # Apply changes
        for change in changes:
            supabase.table(table_name) \
                .update(change['data']) \
                .eq('opgave_id', change['opgave_id']) \
                .execute()

        return True, f"Opdaterede {len(changes)} rækker."
    
    except Exception as e:
        st.exception(e)
        return False, f"Fejl under opdatering: {str(e)}"
    
#Hent næstekommende opgave id
def get_next_opgave_id(df):
    try:
        if not df.empty:
            max_opgave_id = int(df['opgave_id'].max())
        else:
            max_opgave_id = 0
        
        next_opgave_id = max_opgave_id + 1
        return next_opgave_id

    except Exception as e:
        st.error(f"Fejl ved beregning af næste Opgave_id: {str(e)}")
        return 1  # Fallback værdi

# TILFØJ EN RÆKKE I SUPABASE
def append_row(supabase, table_name, row_data):
    try:
        # Ensure all date fields are converted to string format (ISO)
        for key, value in row_data.items():
            if isinstance(value, (datetime.date, pd.Timestamp)):
                row_data[key] = value.isoformat()

        supabase.table(table_name).insert(row_data).execute()
        return True
    except Exception as e:
        st.error(f"Error appending row: {str(e)}")
        return False
    
def handle_deletion(supabase, table_name, edited_data):
    selected_rows = edited_data[edited_data['Select'] == True]

    if not selected_rows.empty:
        st.session_state.confirm_delete = True
        st.session_state.rows_to_delete = selected_rows['opgave_id'].tolist()
    else:
        st.warning("Du skal vælge mindst én række for at slette.")

def confirm_and_execute_deletion(supabase, table_name):
    if st.session_state.get("confirm_delete"):
        st.warning("Er du sikker på, at du vil slette de valgte projekter permanent?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Ja, slet"):
                deleted_count = 0
                for opgave_id in st.session_state.rows_to_delete:
                    try:
                        response = supabase.table(table_name).delete().eq("opgave_id", opgave_id).execute()
                        if response.data is not None:
                            deleted_count += 1
                    except Exception as e:
                        st.error(f"Fejl ved sletning af opgave {opgave_id}: {str(e)}")

                st.success(f"Slettede {deleted_count} projekt(er)")
                st.session_state.confirm_delete = False
                st.session_state.rows_to_delete = []
                st.session_state.data = fetch_data(supabase, table_name)
                st.rerun()
        with col2:
            if st.button("❌ Annuller"):
                st.session_state.confirm_delete = False
                st.session_state.rows_to_delete = []
                st.info("Sletning annulleret.")


# HENT UNIKKE NAVNE FRA KONSULENTER KOLONNEN
def get_unique_names(df, column_name):
    all_names = []
    for names in df[column_name].dropna():
        if isinstance(names, str):
            # Split names and clean whitespace
            all_names.extend([name.strip() for name in names.split(',')])
    return sorted(list(set(all_names)))

#Tabel Export
def export_projects_table(selected_df):
    columns_to_export = [
    "opgave_id", "Opgavetitel", "Opgavetitel_eng", "Status", "Projektnummer", "Kundenavn",
    "English name", "Opgavebeskrivelse", "Opgavebeskrivelse_eng", "Kundebeskrivelse",
    "Client description", "Tidsramme_start", "Tidsramme_slut", "Kontakter", "Opgaveomfang",
    "EVT. TI budgetandel", "Konsulenter", "Land", "Rapport/projektmappe", "Projektpartnere"]
    header_column = 'Opgavetitel'
    
    doc = Document()
    doc.add_heading('Referencer', 0)
    
    for _, row in selected_df.iterrows():
        doc.add_heading(str(row[header_column]), level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        for column in columns_to_export:
            if column != header_column:
                row_cells = table.add_row().cells
                row_cells[0].text = column
                row_cells[1].text = str(row[column])
        
        doc.add_paragraph()
    
    return save_to_bytes(doc)
#Udvalgte data export
def export_projects_Short_presentation(selected_df):
    """Export selected projects to a Word document with a summarized format."""
    doc = Document()
    doc.add_heading('REFERENCER', 0)
    
    for _, row in selected_df.iterrows():
        doc.add_heading(str(row['Opgavetitel']), level=2)
        doc.add_paragraph(row['Opgavebeskrivelse'])
        doc.add_paragraph(f"Kunde: {row['Kundenavn']}")
        doc.add_paragraph(f"Periode: {row['Tidsramme_start']} - {row['Tidsramme_slut']}")
        doc.add_paragraph()
    
    return save_to_bytes(doc)
# Exporter til excel
def export_projects_excel(selected_df):
    output = BytesIO()
    
    # Create Excel file with openpyxl engine
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        selected_df.to_excel(writer, sheet_name='Referencer', index=False)
        
        # Get the openpyxl workbook and worksheet objects
        workbook = writer.book
        worksheet = writer.sheets['Referencer']
        
        # Format headers
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        
        header_font = Font(bold=True)
        header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        header_alignment = Alignment(wrap_text=True, vertical='top')
        thin_border = Border(
            left=Side(style='thin'), 
            right=Side(style='thin'), 
            top=Side(style='thin'), 
            bottom=Side(style='thin')
        )
        
        # Apply formatting to header row
        for col in range(1, len(selected_df.columns) + 1):
            cell = worksheet.cell(row=1, column=col)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        # Auto-adjust column widths
        for col in worksheet.columns:
            max_length = 0
            column = col[0].column_letter
            
            # Find the maximum length in the column
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            # Set the column width
            adjusted_width = (max_length + 2)
            worksheet.column_dimensions[column].width = adjusted_width
    
    output.seek(0)
    return output.getvalue()
#Gem som fil
def save_to_bytes(doc):
    """Save document to BytesIO for download."""
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

#########################################################################
########################## Streamlit frontend ###########################
#########################################################################

def main():
    st.set_page_config(layout="wide", initial_sidebar_state="expanded")
    st.markdown("<h1 style='text-align: center;'>Referencedatabase Erhverv og Samfund 📚</h1>", unsafe_allow_html=True)

    # Initialize Supabase
    supabase = init_supabase()

    # Load data from Supabase
    if "data" not in st.session_state:
        st.session_state.data = fetch_data(supabase, table_name)

    df = st.session_state.data.sort_values(by='opgave_id', ascending=False)
    
   

    # Filters section
    st.write("### Filtre")

    exclude_statuses = ["Udgået som reference", "Afslag", "Ideen er opgivet - projektet slettes fra basen"]
    filtered_df1 = df[~df["Status"].isin(exclude_statuses)]

    filtered_df = filtered_df1.copy()

    unique_names = get_unique_names(filtered_df1, 'Konsulenter')

    # Ensure the columns are datetime
    filtered_df["Tidsramme_start"] = pd.to_datetime(filtered_df["Tidsramme_start"], errors="coerce")
    filtered_df["Tidsramme_slut"] = pd.to_datetime(filtered_df["Tidsramme_slut"], errors="coerce")

    selected_columns = [3, 12, 18]
    cols_per_row = 3

    # Find valid min and max dates, handling NaT values properly
    min_date = pd.concat([filtered_df["Tidsramme_start"], filtered_df["Tidsramme_slut"]]).min()
    max_date = pd.concat([filtered_df["Tidsramme_start"], filtered_df["Tidsramme_slut"]]).max()

    min_date = min_date.date()
    max_date = max_date.date()

    # Create columns for filters
    cols = st.columns(cols_per_row)

    # Date filters
    with cols[0]:
        start_date = st.date_input("Startdato", value=min_date, min_value=min_date, max_value=max_date)
    with cols[1]:
        end_date = st.date_input("Slutdato", value=max_date, min_value=min_date, max_value=max_date)
    with cols[2]:
        selected_consultants = st.multiselect(
            "Filter by Konsulenter",
            options=unique_names
        )

    # Apply consultant filter if any consultants are selected
    if selected_consultants:
        filtered_df = filtered_df[filtered_df['Konsulenter'].apply(
            lambda x: any(consultant in str(x) for consultant in selected_consultants) if isinstance(x, str) else False
        )]

    # Add other filters in the remaining columns
    for i, column_idx in enumerate(selected_columns):
        with cols[(i + 2) % cols_per_row]:  # Start after the date and "Opgaveomfang" filters
            column = filtered_df.columns[column_idx]
            unique_values = sorted(df[column].dropna().unique().astype(str))
            selected_values = st.multiselect(f"Filter by {column}", unique_values, key=f"filter_{column}")
            
            if selected_values:
                filtered_df = filtered_df[filtered_df[column].astype(str).isin(selected_values)]

    # Apply date range filter if both dates are selected - handle NaT values properly
    if start_date and end_date:
        # Convert to pandas Timestamp for comparison
        pd_start_date = pd.Timestamp(start_date)
        pd_end_date = pd.Timestamp(end_date)

        # Filter rows where:
        start_date_mask = (filtered_df["Tidsramme_start"] >= pd_start_date) | (filtered_df["Tidsramme_start"].isna())
        end_date_mask = (filtered_df["Tidsramme_slut"] <= pd_end_date) | (filtered_df["Tidsramme_slut"].isna())

        filtered_df = filtered_df[start_date_mask & end_date_mask]

    # Search functionality
    search_term = st.text_input("Søg i alt tekst", "")
    if search_term:
        mask = filtered_df.astype(str).apply(lambda x: x.str.contains(search_term, case=False, na=False)).any(axis=1)
        filtered_df = filtered_df[mask]

    # Pagination
    rows_per_page = 99
    total_pages = max(1, len(filtered_df) // rows_per_page + (1 if len(filtered_df) % rows_per_page != 0 else 0))
    if "current_page" not in st.session_state:
        st.session_state.current_page = 1
    start_idx = (st.session_state.current_page - 1) * rows_per_page
    end_idx = min(start_idx + rows_per_page, len(filtered_df))

    # First display the table
    st.markdown("<h3 style='text-align: center;'> Referenceprojekter</h3>", unsafe_allow_html=True)
    page_indices = filtered_df.index[start_idx:end_idx]
    
    # Prepare editable columns
    edited_df = filtered_df.loc[page_indices].copy()

    # Add a selection column to the dataframe
    edited_df['Select'] = False
    
    # Reorder columns to put 'Select' first
    cols = ['Select'] + [col for col in edited_df.columns if col != 'Select']
    edited_df = edited_df[cols]

    # Formatering af tabellen, hvor det er muligt at redigere direkte i tabellen
    edited_data = st.data_editor(
        edited_df,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Select": st.column_config.CheckboxColumn("Select")
        }
    )

    # Gem original data før redigering
    if 'original_data' not in st.session_state:
        st.session_state.original_data = filtered_df.copy()

    # Create side-by-side columns for Save and Delete buttons
    col_save, col_delete = st.columns([1, 1])

    with col_save:
        if st.button("💾 Gem ændringer"):
            with st.spinner('Gemmer ændringer...'):
                success, message = update_multiple_rows(
                    supabase,
                    table_name,
                    edited_data,
                    st.session_state.original_data
                )

            if success:
                st.success(message or "Ændringer er gemt")
                st.session_state.data = fetch_data(supabase, table_name)
                st.session_state.original_data = st.session_state.data.copy()
                st.rerun()
            else:
                st.error("Fejl i forsøget på at gemme")
                st.warning(message or "Prøv igen eller kontakt support.")

    with col_delete:
        if st.button("🗑️ Slet valgte projekter"):
            handle_deletion(supabase, table_name, edited_data)

    # Show confirmation if needed
    confirm_and_execute_deletion(supabase, table_name)

    # Skift mellem sider
    col1, col2, col3 = st.columns([1, 30, 1])
    with col1:
        st.button("←", on_click=lambda: setattr(st.session_state, "current_page", max(1, st.session_state.current_page - 1)))
    with col2:
        st.markdown(f"<h3 style='text-align: center;'> Page {st.session_state.current_page} of {total_pages}</h3>", unsafe_allow_html=True)
    with col3:
        st.button("→", on_click=lambda: setattr(st.session_state, "current_page", min(total_pages, st.session_state.current_page + 1)))

    st.markdown("---")

    # Tilføj nyt projekt vha. Form
    with st.form("new_row_form"):
        st.markdown("<h3 style='text-align: center;'> Tilføj nyt projekt i formularen </h3>", unsafe_allow_html=True)
        st.markdown("<h8 style='text-align: left;'> Hold musen over spørgsmålstegnene for at få guides og hjælp til at udfylde felterne  </h8>", unsafe_allow_html=True)
        cols = st.columns(4)

        # Add form fields for your columns
        new_row_data = {}
        with cols[0]:
            next_id = get_next_opgave_id(df)
            existing_statuses = df["Status"].dropna().unique().tolist()
            new_row_data["opgave_id"] = st.text_input(
                "Opgave id",
                value=str(next_id),
                disabled=False,
                help=f"Automatisk genereret ID. Næste ledige nummer: {next_id}"
            )
            new_row_data["Opgavetitel"] = st.text_area("Projekttitel", height=68)
            new_row_data["Opgavetitel_eng"] = st.text_area("Projekttitel engelsk", height=68)
            new_row_data["Status"] = st.selectbox("Status", ['Vælg status'] + existing_statuses)
            new_row_data["Projektnummer"] = st.text_area("Evt. projektnummer", height=68)
            
        with cols[1]:
            new_row_data["Kundenavn"] = st.text_area("Kundenavn", height=68)
            new_row_data["English name"] = st.text_area("Kundenavn engelsk", height=68)
            new_row_data["Opgavebeskrivelse"] = st.text_area(
                "Beskrivelse",
                help="**Husk at skrive i datid og inkluder metode i Inholds-afsnitet**",
                value="Formål \n\nIndhold \n\nResultater",
                height=150
            )
            new_row_data["Opgavebeskrivelse_eng"] = st.text_area(
                "Beskrivelse engelsk",
                help="**Remember to write in past tense, and include methodology in the content section**",
                value="Purpose \n\nContent \n\nResults",
                height=150
            )
            
        with cols[2]:
            new_row_data["Kundebeskrivelse"] = st.text_area("Kundebeskrivelse", height=68)
            new_row_data["Client description"] = st.text_area("Kundebeskrivelse engelsk", height=68)
            new_row_data["Tidsramme_start"] = st.date_input("Start Dato")
            new_row_data["Tidsramme_slut"] = st.date_input("Slut Dato")
            new_row_data["Konsulenter"] = st.text_area("Konsulenter(Projektleder)", placeholder="Skriv fuldt navn på konsulent(er)", height=68)
            new_row_data["EVT. TI budgetandel"] = st.text_area("Budgetandel", placeholder="Tilføj TI's del af budgettet", height=68)

        with cols[3]:
            new_row_data["Opgaveomfang"] = st.text_area("Opgaveomfang", placeholder="Beløbet skrives i DKK", height=68)
            new_row_data["Land"] = st.text_area("Land", height=68)
            new_row_data["Rapport/projektmappe"] = st.text_area("Rapport/projektmappe", placeholder="Indsæt fil-sti eller link", height=68)
            new_row_data["Projektpartnere"] = st.text_area("Projektpartnere", placeholder= "Skriv samarbejdsorganisationer ind", height=68)
            new_row_data["Kontakter"] = st.text_area("Kontakter", placeholder= "Skriv gerne\nKontaktperson \nMail \nTlf", height= 120)

        # Submit button
        submitted = st.form_submit_button("**Tilføj Projekt**")
        if submitted:
            # Validate required fields
            required_fields = ["Opgavetitel", "Status", "Kundenavn"]
            missing_fields = [field for field in required_fields if not new_row_data.get(field)]
            
            if missing_fields:
                st.error(f"Følgende felter skal udfyldes: {', '.join(missing_fields)}")
            else:  
                if append_row(supabase, table_name, new_row_data):
                    st.success("Nyt projekt tilføjet!")
                    # Clear the form inputs by resetting session state
                    for key in list(st.session_state.keys()):
                        if key.startswith("new_"):
                            del st.session_state[key]
                            
                    # Refresh data
                    st.session_state.data = fetch_data(supabase, table_name)
                    st.rerun()

    st.markdown("---")
    st.write("Ved fejl, forbedringsforslag eller andet, skriv gerne til antg@teknologisk.dk")

    #########################################
    ################ Side-bar ###############
    #########################################

    with st.sidebar:
        fil_navn = f"Eksporterede_referencer_{datetime.datetime.now().strftime('%Y-%m-%d')}"
        
        st.markdown("<h1 style='text-align: center;'> AI hjælp </h1>", unsafe_allow_html=True)
        st.write("Få hjælp til at finde referenceprojekter af en AI-assistent, skal du klikke på knappen nednunder ⬇️. Herefter skal du vælge assistenten: Asistent til Referencedatabase. Fra menuen til højre på siden. Sti: Home / Organization / C016")
        st.link_button("Gå til AI-assistent 🤖",'https://ai.localdom.net/#!/')
        st.markdown("---")

        # Eksporter projekter
        st.markdown("<h1 style='text-align: center;'> Eksporter </h1>", unsafe_allow_html=True)
        st.write("Marker et eller flere projekter i tabellen. Herefter kan du eksportere projekterne til et Word dokument eller Excel-ark.")

        format_choice = st.radio("Vælg eksportformat:", ["Tabelformat - Word", "Excel - format"])

        if st.button("Eksporter valgte projekter"):
            selected_df = edited_data[edited_data['Select'] == True].drop('Select', axis=1)
            
            if len(selected_df) > 0:
                if format_choice == "Tabelformat - Word":
                    doc_io = export_projects_table(selected_df)
                    st.download_button(
                        label="Download Word Dokument",
                        data=doc_io.getvalue(),
                        file_name=f"{fil_navn}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    excel_data = export_projects_excel(selected_df)
                    st.download_button(
                        label="Download Excel",
                        data=excel_data,
                        file_name=f"{fil_navn}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
            else:
                st.warning("Vælg mindst ét projekt at eksportere")

if __name__ == "__main__":
    main()
